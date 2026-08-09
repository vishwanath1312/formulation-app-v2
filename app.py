import os
import warnings
import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler, PolynomialFeatures
from sklearn.linear_model import LinearRegression, Ridge
from sklearn.tree import DecisionTreeRegressor
from sklearn.ensemble import RandomForestRegressor
from sklearn.svm import SVR
from xgboost import XGBRegressor
from sklearn.multioutput import MultiOutputRegressor
from sklearn.model_selection import GridSearchCV, LeaveOneOut, cross_val_predict
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error
import joblib
import statsmodels.api as sm
import statsmodels.formula.api as smf
from scipy.stats import zscore
from scipy import stats as scipy_stats
from scipy.optimize import differential_evolution
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore")  # small-sample sklearn convergence/feature-name warnings are just noise here

st.set_page_config(page_title="Formulation Lab | DoE Dashboard", page_icon="🧬", layout="wide")

# ----------------------------------------------------------------------
# Design system
# ----------------------------------------------------------------------
# A "lab instrument" identity for a DoE dashboard: cool near-white canvas
# (not the usual warm-cream AI-tool default), a single teal accent, and
# monospace numerals so data reads like a measurement readout rather than
# marketing copy. Tokens:
#   canvas #F5F7F8   panel #FFFFFF   ink #101820   slate #55606B
#   teal (primary)   #0E6E62        amber (highlight) #C97A2B
# Display face: Space Grotesk · Body: Inter · Data: IBM Plex Mono
CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@500;600;700&family=Inter:wght@400;500;600&family=IBM+Plex+Mono:wght@400;500;600&display=swap');

:root {
  --canvas: #F5F7F8;
  --panel: #FFFFFF;
  --ink: #101820;
  --slate: #55606B;
  --teal: #0E6E62;
  --teal-dark: #0B5850;
  --amber: #C97A2B;
  --border: #E2E6EA;
}

.stApp { background-color: var(--canvas); }

h1, h2, h3, h4, .stApp [data-testid="stMarkdownContainer"] h1,
.stApp [data-testid="stMarkdownContainer"] h2, .stApp [data-testid="stMarkdownContainer"] h3 {
  font-family: 'Space Grotesk', sans-serif !important;
  color: var(--ink) !important;
  letter-spacing: -0.01em;
}
html, body, [class*="css"], .stApp, p, label, .stMarkdown {
  font-family: 'Inter', sans-serif;
}
.stApp [data-testid="stMetricValue"], code, .stDataFrame, [data-testid="stDataFrame"] * {
  font-family: 'IBM Plex Mono', monospace !important;
}

/* Hero header */
.app-eyebrow {
  font-family: 'IBM Plex Mono', monospace;
  font-size: 0.72rem;
  letter-spacing: 0.16em;
  text-transform: uppercase;
  color: var(--teal);
  margin-bottom: 0.3rem;
  font-weight: 600;
}
.app-hero-title {
  font-family: 'Space Grotesk', sans-serif;
  font-weight: 700;
  font-size: 2.1rem;
  color: var(--ink);
  margin: 0 0 0.3rem 0;
  line-height: 1.15;
}
.app-hero-sub { color: var(--slate); font-size: 0.96rem; margin-bottom: 0.7rem; max-width: 70ch; }
.app-hr { border: none; border-top: 1px solid var(--border); margin: 0 0 1.3rem 0; }

/* Spec strip: a mono "readout" line of key facts under a header */
.spec-strip {
  display: flex; gap: 1.6rem; flex-wrap: wrap;
  font-family: 'IBM Plex Mono', monospace; font-size: 0.76rem; color: var(--slate);
  border-top: 1px solid var(--border); border-bottom: 1px solid var(--border);
  padding: 0.55rem 0.1rem; margin: -0.6rem 0 1.5rem 0;
}
.spec-strip b { color: var(--ink); font-weight: 600; }

/* Sidebar as a dark console panel */
[data-testid="stSidebar"] { background-color: var(--ink); }
[data-testid="stSidebar"] * { color: #E8EDF0 !important; }
[data-testid="stSidebar"] .app-eyebrow { color: #57C7B8 !important; }
[data-testid="stSidebar"] hr { border-color: rgba(255,255,255,0.14); }
[data-testid="stSidebar"] [data-testid="stMetric"] { background: rgba(255,255,255,0.06); border-color: rgba(255,255,255,0.14); }

/* Metric / KPI cards */
[data-testid="stMetric"] {
  background: var(--panel);
  border: 1px solid var(--border);
  border-radius: 10px;
  padding: 0.85rem 1rem;
}

/* Buttons */
.stButton>button, .stDownloadButton>button {
  background-color: var(--teal); color: #fff !important; border: none !important;
  border-radius: 8px !important; font-weight: 600; padding: 0.5rem 1.1rem;
}
.stButton>button:hover, .stDownloadButton>button:hover { background-color: var(--teal-dark); }

/* Dataframes / tables */
[data-testid="stDataFrame"] { border: 1px solid var(--border); border-radius: 10px; overflow: hidden; }

/* Tabs & selects: tidy borders */
[data-baseweb="select"] > div { border-radius: 8px !important; border-color: var(--border) !important; }
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

TEAL_SCALE = [[0.0, "#F5F7F8"], [0.35, "#9FD6CB"], [0.7, "#33978A"], [1.0, "#0E6E62"]]


def render_header(eyebrow, title, subtitle=None, facts=None):
    """Consistent hero header used at the top of every page."""
    st.markdown(f'<div class="app-eyebrow">{eyebrow}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="app-hero-title">{title}</div>', unsafe_allow_html=True)
    if subtitle:
        st.markdown(f'<div class="app-hero-sub">{subtitle}</div>', unsafe_allow_html=True)
    if facts:
        strip = "".join(f"<span>{f}</span>" for f in facts)
        st.markdown(f'<div class="spec-strip">{strip}</div>', unsafe_allow_html=True)
    else:
        st.markdown('<hr class="app-hr"/>', unsafe_allow_html=True)


# Ensure the folder for saved models exists (needed before joblib.dump)
os.makedirs("models", exist_ok=True)

DATA_PATH = os.path.join("data", "data of the formulation.xlsx")


# ----------------------------------------------------------------------
# Load dataset
# ----------------------------------------------------------------------
@st.cache_data
def load_data():
    return pd.read_excel(DATA_PATH)


data = load_data()
X = data[["Stearic acid", "Tween 80"]]
y = data[["Entrapment efficiency", "Drug content", "Drug release", "Particle size"]]


# ----------------------------------------------------------------------
# Journal-ready RSM ANOVA (Design-Expert-style quadratic model)
# ----------------------------------------------------------------------
# Fits the standard second-order response-surface model
#   y = b0 + b1*A + b2*B + b12*AB + b11*A^2 + b22*B^2
# and decomposes it into the ANOVA table format expected by formulation /
# pharmaceutics journals (Model, each term, Residual, Lack of Fit, Pure
# Error, Cor Total), plus the fit-quality statistics (R^2, Adjusted R^2,
# Predicted R^2, Adequate Precision, %CV) that reviewers expect to see
# alongside it. This intentionally uses plain (unregularized) OLS, since
# that is the convention every RSM/DoE paper and Design-Expert itself
# follows for this table — the ⭐ prediction models elsewhere in the app
# remain the tuned/regularized pipelines used for actual predictions.
RSM_TERM_LABELS = {
    "A": "A — Stearic acid", "B": "B — Tween 80", "AB": "AB",
    "A2": "A²", "B2": "B²",
}


@st.cache_data
def compute_rsm_anova(data_in, response, x1="Stearic acid", x2="Tween 80"):
    df = data_in.rename(columns={x1: "A", x2: "B", response: "y"})
    df["AB"] = df["A"] * df["B"]
    df["A2"] = df["A"] ** 2
    df["B2"] = df["B"] ** 2

    n = len(df)
    model = smf.ols("y ~ A + B + AB + A2 + B2", data=df).fit()
    p_params = int(model.df_model) + 1
    resid_df = int(model.df_resid)
    model_df = int(model.df_model)

    ss_total = float(np.sum((df["y"] - df["y"].mean()) ** 2))
    ss_resid = float(np.sum(model.resid ** 2))
    ss_model = ss_total - ss_resid
    ms_model = ss_model / model_df
    ms_resid = ss_resid / resid_df
    f_model = ms_model / ms_resid
    p_model = float(1 - scipy_stats.f.cdf(f_model, model_df, resid_df))

    # Lack of fit / pure error, using any replicated (A, B) design points
    pure_error_ss, pure_error_df = 0.0, 0
    for _, g in df.groupby(["A", "B"])["y"]:
        if len(g) > 1:
            pure_error_ss += float(np.sum((g - g.mean()) ** 2))
            pure_error_df += len(g) - 1
    lof_ss = ss_resid - pure_error_ss
    lof_df = resid_df - pure_error_df
    has_replicates = pure_error_df > 0
    ms_lof = lof_ss / lof_df if lof_df > 0 else np.nan
    ms_pe = pure_error_ss / pure_error_df if has_replicates else np.nan
    f_lof = ms_lof / ms_pe if has_replicates and lof_df > 0 else np.nan
    p_lof = float(1 - scipy_stats.f.cdf(f_lof, lof_df, pure_error_df)) if has_replicates and lof_df > 0 else np.nan

    def sig_label(pval):
        if pval is None or (isinstance(pval, float) and np.isnan(pval)):
            return ""
        return "significant" if pval < 0.05 else "not significant"

    rows = [{"Source": "Model", "SS": ss_model, "df": model_df, "MS": ms_model,
             "F-value": f_model, "p-value (Prob > F)": p_model, "": sig_label(p_model)}]
    for term in ["A", "B", "AB", "A2", "B2"]:
        t = model.tvalues[term]
        pval = float(model.pvalues[term])
        f_term = t ** 2
        ss_term = f_term * ms_resid
        rows.append({"Source": RSM_TERM_LABELS[term], "SS": ss_term, "df": 1, "MS": ss_term,
                     "F-value": f_term, "p-value (Prob > F)": pval, "": sig_label(pval)})
    rows.append({"Source": "Residual", "SS": ss_resid, "df": resid_df, "MS": ms_resid,
                 "F-value": np.nan, "p-value (Prob > F)": np.nan, "": ""})
    rows.append({"Source": "Lack of Fit", "SS": lof_ss, "df": lof_df, "MS": ms_lof,
                 "F-value": f_lof, "p-value (Prob > F)": p_lof,
                 "": sig_label(p_lof) if has_replicates else "n/a (no replicates)"})
    rows.append({"Source": "Pure Error", "SS": pure_error_ss, "df": pure_error_df, "MS": ms_pe,
                 "F-value": np.nan, "p-value (Prob > F)": np.nan, "": ""})
    rows.append({"Source": "Cor Total", "SS": ss_total, "df": n - 1, "MS": np.nan,
                 "F-value": np.nan, "p-value (Prob > F)": np.nan, "": ""})
    anova_df = pd.DataFrame(rows)

    # Predicted R^2 via PRESS (leave-one-out deletion residuals from the hat matrix)
    hat = model.get_influence().hat_matrix_diag
    press = float(np.sum((model.resid / (1 - hat)) ** 2))
    pred_r2 = 1 - press / ss_total

    adeq_precision = (model.fittedvalues.max() - model.fittedvalues.min()) / np.sqrt(p_params * ms_resid / n)
    cv_pct = np.sqrt(ms_resid) / df["y"].mean() * 100

    fit_stats = {
        "R²": model.rsquared, "Adjusted R²": model.rsquared_adj, "Predicted R²": pred_r2,
        "Adequate Precision": adeq_precision, "C.V. %": cv_pct,
        "Std. Dev.": np.sqrt(ms_resid), "Mean": df["y"].mean(), "PRESS": press,
    }

    # Equation in actual factors (raw units)
    c = model.params
    eq_actual = (
        f"{response} = {c['Intercept']:+.6g} {c['A']:+.6g}·A {c['B']:+.6g}·B "
        f"{c['AB']:+.6g}·AB {c['A2']:+.6g}·A² {c['B2']:+.6g}·B²"
    )

    # Equation in coded factors (-1 / +1 over the explored range) — coefficient
    # magnitude is then directly comparable across terms, as journals show it.
    df_coded = df.copy()
    for col in ["A", "B"]:
        lo, hi = df[col].min(), df[col].max()
        mid, half = (hi + lo) / 2, (hi - lo) / 2
        df_coded[col] = (df[col] - mid) / half
    df_coded["AB"] = df_coded["A"] * df_coded["B"]
    df_coded["A2"] = df_coded["A"] ** 2
    df_coded["B2"] = df_coded["B"] ** 2
    model_coded = smf.ols("y ~ A + B + AB + A2 + B2", data=df_coded).fit()
    cc = model_coded.params
    eq_coded = (
        f"{response} = {cc['Intercept']:+.6g} {cc['A']:+.6g}·A {cc['B']:+.6g}·B "
        f"{cc['AB']:+.6g}·AB {cc['A2']:+.6g}·A² {cc['B2']:+.6g}·B²"
    )

    # Multicollinearity check (condition number on the coded design matrix)
    cond_no = float(model_coded.condition_number)

    return {
        "anova_df": anova_df, "fit_stats": fit_stats, "eq_actual": eq_actual,
        "eq_coded": eq_coded, "cond_no": cond_no, "p_model": p_model, "p_lof": p_lof,
        "has_replicates": has_replicates, "n": n, "model_df": model_df, "resid_df": resid_df,
    }


def _docx_set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def build_anova_docx(response, result):
    """Builds a journal-style Table (ANOVA + fit statistics + equation) as a .docx in memory."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run(f"Table. ANOVA for Response Surface Quadratic Model — {response}")
    run.bold = True
    run.font.size = Pt(11)

    anova_df = result["anova_df"]
    n_rows, n_cols = anova_df.shape[0] + 1, anova_df.shape[1]
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = list(anova_df.columns)
    headers[-1] = "Significance"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_set_cell_shading(cell, "D6EFE9")

    for i, row in anova_df.iterrows():
        for j, col in enumerate(anova_df.columns):
            val = row[col]
            if col in ("SS", "MS"):
                text = "" if pd.isna(val) else f"{val:.4f}"
            elif col == "F-value":
                text = "" if pd.isna(val) else f"{val:.2f}"
            elif col == "p-value (Prob > F)":
                text = "" if pd.isna(val) else ("< 0.0001" if val < 0.0001 else f"{val:.4f}")
            elif col == "df":
                text = "" if pd.isna(val) else f"{int(val)}"
            else:
                text = str(val)
            cell = table.rows[i + 1].cells[j]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if row["Source"] in ("Model", "Cor Total"):
                for r in cell.paragraphs[0].runs:
                    r.bold = True

    doc.add_paragraph()
    fs = result["fit_stats"]
    fs_para = doc.add_paragraph()
    fs_para.add_run(
        f"R² = {fs['R²']:.4f}; Adjusted R² = {fs['Adjusted R²']:.4f}; "
        f"Predicted R² = {fs['Predicted R²']:.4f}; Adequate Precision = "
        f"{fs['Adequate Precision']:.3f}; C.V.% = {fs['C.V. %']:.2f}; "
        f"Std. Dev. = {fs['Std. Dev.']:.4f}."
    ).italic = True

    doc.add_paragraph()
    eq_para = doc.add_paragraph()
    eq_para.add_run("Final Equation in Terms of Actual Factors:").bold = True
    doc.add_paragraph(result["eq_actual"])
    eq_para2 = doc.add_paragraph()
    eq_para2.add_run("Final Equation in Terms of Coded Factors:").bold = True
    doc.add_paragraph(result["eq_coded"])

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ----------------------------------------------------------------------
# Model definitions
# ----------------------------------------------------------------------
# Every model is a Pipeline (StandardScaler + estimator) so all six models
# share one calling convention: models[name].predict([[x1, x2]]).
# Hyperparameters are tuned via leave-one-out cross-validation (LOOCV) —
# with only 10 experimental runs, a single train/test split is far too
# noisy to trust, so every fold gets to act as the test set exactly once.
def build_model_specs():
    return {
        "Linear Regression": (
            Pipeline([("scaler", StandardScaler()), ("reg", MultiOutputRegressor(LinearRegression()))]),
            {},
        ),
        "Polynomial Regression (RSM, Ridge)": (
            # Ridge (L2-regularized) instead of plain LinearRegression on the
            # degree-2 features: the raw quadratic RSM design is highly
            # collinear (confirmed by the large condition number on the
            # ANOVA page), and regularization stabilizes those coefficients.
            Pipeline([
                ("poly", PolynomialFeatures(degree=2)),
                ("scaler", StandardScaler()),
                ("reg", MultiOutputRegressor(Ridge())),
            ]),
            {"reg__estimator__alpha": [0.01, 0.1, 1, 10, 100]},
        ),
        "Decision Tree": (
            Pipeline([("scaler", StandardScaler()), ("reg", MultiOutputRegressor(DecisionTreeRegressor(random_state=42)))]),
            {"reg__estimator__max_depth": [2, 3, None], "reg__estimator__min_samples_leaf": [1, 2]},
        ),
        "Random Forest": (
            Pipeline([("scaler", StandardScaler()), ("reg", MultiOutputRegressor(RandomForestRegressor(n_estimators=100, random_state=42)))]),
            {"reg__estimator__max_depth": [2, 3], "reg__estimator__min_samples_leaf": [1, 2]},
        ),
        "SVR": (
            Pipeline([("scaler", StandardScaler()), ("reg", MultiOutputRegressor(SVR(kernel="rbf")))]),
            {"reg__estimator__C": [0.1, 1, 10], "reg__estimator__gamma": ["scale", 0.1],
             "reg__estimator__epsilon": [0.01, 0.1]},
        ),
        "XGBoost": (
            Pipeline([("scaler", StandardScaler()), ("reg", MultiOutputRegressor(XGBRegressor(objective="reg:squarederror", random_state=42)))]),
            {"reg__estimator__n_estimators": [50, 100], "reg__estimator__max_depth": [2, 3],
             "reg__estimator__learning_rate": [0.1, 0.2]},
        ),
    }


# ----------------------------------------------------------------------
# Train + tune models (cached so this only runs once per session)
# ----------------------------------------------------------------------
@st.cache_resource
def train_models(X, y):
    loo = LeaveOneOut()
    model_specs = build_model_specs()

    models = {}
    best_params = {}
    metrics = {}

    for name, (pipe, grid) in model_specs.items():
        if grid:
            search = GridSearchCV(pipe, grid, cv=loo, scoring="r2", n_jobs=1)
            search.fit(X, y)
            best_est = search.best_estimator_
            best_params[name] = search.best_params_
        else:
            best_est = pipe.fit(X, y)
            best_params[name] = {}

        # Out-of-fold (leave-one-out) predictions using the tuned hyperparameters
        # give an honest estimate of how the model performs on unseen runs,
        # using every one of the 10 experimental points as a held-out test case.
        oof_pred = cross_val_predict(best_est, X, y, cv=loo, n_jobs=1)
        metrics[name] = {
            "R² (LOOCV)": r2_score(y, oof_pred),
            "MAE (LOOCV)": mean_absolute_error(y, oof_pred),
            "MSE (LOOCV)": mean_squared_error(y, oof_pred),
            "RMSE (LOOCV)": np.sqrt(mean_squared_error(y, oof_pred)),
        }

        models[name] = best_est  # already refit on ALL data by GridSearchCV(refit=True)
        joblib.dump(best_est, f"models/{name.replace(' ', '_').replace(',', '').replace('(', '').replace(')', '').lower()}.pkl")

    best_model_name = max(metrics, key=lambda n: metrics[n]["R² (LOOCV)"])
    return models, best_params, metrics, best_model_name


with st.spinner("Training and tuning models with leave-one-out cross-validation..."):
    models, best_params, metrics, best_model_name = train_models(X, y)


# ----------------------------------------------------------------------
# Multi-page dashboard
# ----------------------------------------------------------------------
st.sidebar.markdown('<div class="app-eyebrow">Lipid Nanoparticle DoE</div>', unsafe_allow_html=True)
st.sidebar.markdown(
    '<div style="font-family:\'Space Grotesk\',sans-serif;font-weight:700;'
    'font-size:1.35rem;margin-bottom:1rem;">🧬 Formulation Lab</div>',
    unsafe_allow_html=True,
)
page = st.sidebar.radio(
    "Go to",
    ["Dataset", "Prediction", "Reverse Prediction", "Model Comparison", "ANOVA Analysis",
     "Response Surfaces", "Optimization", "Outlier Analysis"],
)
st.sidebar.markdown("---")
st.sidebar.markdown(
    '<div style="background:rgba(255,255,255,0.06);border:1px solid rgba(255,255,255,0.14);'
    'border-radius:10px;padding:0.8rem 0.9rem;">'
    '<div class="app-eyebrow" style="margin-bottom:0.4rem;">Best model (LOOCV R²)</div>'
    f'<div style="font-family:\'Space Grotesk\',sans-serif;font-weight:600;font-size:1.05rem;'
    f'line-height:1.25;">{best_model_name}</div>'
    f'<div style="font-family:\'IBM Plex Mono\',monospace;font-size:0.82rem;color:#57C7B8;'
    f'margin-top:0.3rem;">↑ R² = {metrics[best_model_name]["R² (LOOCV)"]:.3f}</div>'
    '</div>',
    unsafe_allow_html=True,
)

# ---------------- Dataset ----------------
if page == "Dataset":
    render_header(
        "Raw Data",
        "Original Dataset",
        "The experimental design matrix used to train and validate every model in this app.",
        facts=[f"<b>{data.shape[0]}</b> runs", f"<b>{data.shape[1]}</b> columns",
               "<b>2</b> factors", "<b>4</b> responses"],
    )
    st.dataframe(data, width='stretch', hide_index=True)

    st.markdown("#### Summary Statistics")
    st.dataframe(data.describe().T, width='stretch')

    st.download_button(
        "Download dataset as CSV",
        data.to_csv(index=False),
        file_name="data_of_the_formulation.csv",
        mime="text/csv",
    )

# ---------------- Prediction UI ----------------
elif page == "Prediction":
    render_header(
        "Forward Model",
        "Formulation Predictor",
        "Enter a candidate formulation and compare predictions across all six tuned models.",
    )
    st.sidebar.header("Input Parameters")
    stearic = st.sidebar.number_input("Stearic acid", min_value=60, max_value=400, step=10, value=240)
    tween = st.sidebar.number_input("Tween 80", min_value=60, max_value=200, step=10, value=120)

    if st.sidebar.button("Predict"):
        output_cols = ["Entrapment efficiency", "Drug content", "Drug release", "Particle size"]
        rows, index = [], []
        for name, model in models.items():
            rows.append(model.predict([[stearic, tween]])[0])
            index.append(f"⭐ {name}" if name == best_model_name else name)

        st.write("#### Predictions from all models")
        st.table(pd.DataFrame(rows, index=index, columns=output_cols))
        st.caption(
            f"⭐ = model with the best leave-one-out cross-validated R² "
            f"({best_model_name}, R² = {metrics[best_model_name]['R² (LOOCV)']:.3f}). "
            "See the Model Comparison page for full details."
        )

# ---------------- Reverse Prediction (Inverse Design) ----------------
elif page == "Reverse Prediction":
    render_header(
        "Inverse Design",
        "Reverse Prediction",
        "Specify the desired outcome and search the design space for the "
        "Stearic acid / Tween 80 combination that best achieves it.",
    )
    st.caption(
        "This is an inverse (many-to-few) problem: 4 targets, 2 tunable "
        "inputs, so an exact match usually isn't possible. The search finds "
        "the best achievable compromise (minimum weighted error) within the "
        "explored design space."
    )

    forward_model_name = st.selectbox(
        "Forward model to search against",
        list(models.keys()),
        index=list(models.keys()).index(best_model_name),
        help="The reverse search evaluates candidate (Stearic acid, Tween 80) "
             "pairs using this forward model's predictions. Defaults to the "
             "model with the best leave-one-out cross-validated R².",
    )

    st.markdown("### Target Responses")
    c1, c2, c3, c4 = st.columns(4)
    target_ee = c1.number_input(
        "Target Entrapment efficiency", value=float(data["Entrapment efficiency"].mean()),
        step=0.01, format="%.4f")
    target_dc = c2.number_input(
        "Target Drug content", value=float(data["Drug content"].mean()),
        step=0.01, format="%.4f")
    target_dr = c3.number_input(
        "Target Drug release", value=float(data["Drug release"].mean()),
        step=0.01, format="%.4f")
    target_ps = c4.number_input(
        "Target Particle size", value=float(data["Particle size"].mean()), step=1.0)

    st.markdown("### Response Weights (optional)")
    st.caption("Increase a weight to prioritize matching that response more closely.")
    w1, w2, w3, w4 = st.columns(4)
    weight_ee = w1.slider("Entrapment efficiency weight", 0.0, 5.0, 1.0, step=0.5)
    weight_dc = w2.slider("Drug content weight", 0.0, 5.0, 1.0, step=0.5)
    weight_dr = w3.slider("Drug release weight", 0.0, 5.0, 1.0, step=0.5)
    weight_ps = w4.slider("Particle size weight", 0.0, 5.0, 1.0, step=0.5)

    if st.button("Find Matching Formulation"):
        targets = np.array([target_ee, target_dc, target_dr, target_ps])
        weights = np.array([weight_ee, weight_dc, weight_dr, weight_ps])
        y_ranges = (y.max() - y.min()).to_numpy()
        y_ranges[y_ranges == 0] = 1.0  # avoid div-by-zero

        def forward_predict(x1, x2):
            return models[forward_model_name].predict([[x1, x2]])[0]

        def objective(params):
            x1, x2 = params
            pred = forward_predict(x1, x2)
            diff = (pred - targets) / y_ranges
            return np.sum(weights * diff ** 2)

        bounds = [
            (float(X["Stearic acid"].min()), float(X["Stearic acid"].max())),
            (float(X["Tween 80"].min()), float(X["Tween 80"].max())),
        ]
        with st.spinner("Searching the design space..."):
            result = differential_evolution(objective, bounds, seed=42, tol=1e-10, polish=True)
        x1_opt, x2_opt = result.x
        pred_opt = forward_predict(x1_opt, x2_opt)

        st.success("Best-matching formulation found", icon="✅")
        c1, c2 = st.columns(2)
        c1.metric("Stearic acid (mg)", f"{x1_opt:.2f}")
        c2.metric("Tween 80 (mg)", f"{x2_opt:.2f}")

        compare_df = pd.DataFrame({
            "Response": y.columns.tolist(),
            "Target": targets,
            "Predicted at solution": np.round(pred_opt, 4),
            "Abs. difference": np.round(np.abs(pred_opt - targets), 4),
        })
        st.markdown("### Target vs. Achieved")
        st.table(compare_df)

        st.download_button(
            "Download result as CSV",
            pd.DataFrame({
                "Stearic acid": [x1_opt], "Tween 80": [x2_opt],
                "Predicted Entrapment efficiency": [pred_opt[0]],
                "Predicted Drug content": [pred_opt[1]],
                "Predicted Drug release": [pred_opt[2]],
                "Predicted Particle size": [pred_opt[3]],
            }).to_csv(index=False),
            file_name="reverse_prediction_result.csv",
            mime="text/csv",
        )
        st.caption(
            "Note: the search is bounded to the experimentally explored "
            "range (Stearic acid "
            f"{X['Stearic acid'].min():.0f}–{X['Stearic acid'].max():.0f} mg, "
            f"Tween 80 {X['Tween 80'].min():.0f}–{X['Tween 80'].max():.0f} mg) "
            "so results stay within the validated design space."
        )

# ---------------- Model comparison ----------------
elif page == "Model Comparison":
    render_header(
        "Validation",
        "Model Comparison",
        "Leave-one-out cross-validated performance across all six tuned models.",
    )
    st.markdown(
        "Each of the 10 experimental runs is held out and predicted "
        "exactly once by a model trained on the other 9. This is far more "
        "reliable than a single random train/test split on a 10-row dataset, "
        "where the test set would only be 1–2 points."
    )

    metrics_df = pd.DataFrame(metrics).T
    metrics_df = metrics_df.sort_values("R² (LOOCV)", ascending=False)
    metrics_df = metrics_df.reset_index().rename(columns={"index": "Model"})

    def highlight_best(row):
        return ["background-color: #D6EFE9" if row["Model"] == best_model_name else "" for _ in row]

    st.dataframe(
        metrics_df.style.apply(highlight_best, axis=1).format(
            {c: "{:.4f}" for c in metrics_df.columns if c != "Model"}
        ),
        width='stretch',
        hide_index=True,
        column_config={"Model": st.column_config.TextColumn("Model", width="medium")},
    )
    st.success(
        f"Best model: **{best_model_name}** "
        f"(R² = {metrics[best_model_name]['R² (LOOCV)']:.3f}). "
        "This is the model used as the default in Reverse Prediction and "
        "the ⭐-marked row in Prediction.",
        icon="🏆",
    )

    st.markdown("#### Tuned Hyperparameters")
    st.caption(
        "Hyperparameters were selected via GridSearchCV using the same "
        "leave-one-out cross-validation splits, favoring settings that "
        "generalize rather than settings that just fit the training data best."
    )
    for name, params in best_params.items():
        if params:
            st.write(f"**{name}:** {params}")
        else:
            st.write(f"**{name}:** (no tunable hyperparameters)")

    st.caption(
        "Even with tuning, a 10-run dataset limits how accurate any model can "
        "be — the regularized Polynomial (RSM) model tends to generalize best "
        "here because it matches the underlying 2-factor design-of-experiments "
        "structure, while more flexible models (Random Forest, SVR, XGBoost, "
        "Decision Tree) have too little data to reliably learn complex patterns."
    )

# ---------------- ANOVA ----------------
elif page == "ANOVA Analysis":
    render_header(
        "Response Surface Methodology",
        "ANOVA for the Quadratic Model",
        "Design-Expert-style analysis of variance, ready to report in a journal "
        "Results section: significance of the quadratic model and each term, "
        "lack-of-fit test, fit statistics, and the final regression equation.",
    )
    response_name = st.selectbox("Response", y.columns.tolist())
    result = compute_rsm_anova(data, response_name)
    anova_df = result["anova_df"]
    fs = result["fit_stats"]

    st.caption(
        "Fitted model: y = b₀ + b₁A + b₂B + b₁₂AB + b₁₁A² + b₂₂B², where A = "
        "Stearic acid and B = Tween 80. This is plain (unregularized) OLS — the "
        "convention this table follows in RSM/DoE papers — and is separate from "
        "the tuned/regularized models used elsewhere in the app for prediction."
    )

    p_model_str = "p < 0.0001" if result["p_model"] < 0.0001 else f"p = {result['p_model']:.4f}"
    if result["p_model"] < 0.05:
        model_f_value = anova_df.loc[anova_df["Source"] == "Model", "F-value"].iloc[0]
        st.success(
            f"The Model F-value of {model_f_value:.2f} implies the model is "
            f"significant ({p_model_str}). There is only a small chance that an "
            "F-value this large could occur due to noise.",
            icon="✅",
        )
    else:
        st.warning(
            f"The Model F-value implies the model is **not significant** "
            f"(p = {result['p_model']:.4f}) — treat this response's equation as "
            "exploratory rather than a validated predictive model.",
            icon="⚠️",
        )

    if result["has_replicates"]:
        if not np.isnan(result["p_lof"]) and result["p_lof"] >= 0.05:
            st.info(
                f"Lack of Fit p-value = {result['p_lof']:.4f} (not significant) — "
                "this is desirable, meaning the quadratic model fits the data well "
                "relative to the pure error from replicate runs.",
                icon="ℹ️",
            )
        elif not np.isnan(result["p_lof"]):
            st.warning(
                f"Lack of Fit p-value = {result['p_lof']:.4f} (significant) — the "
                "model may not adequately capture the true response surface for "
                "this response.",
                icon="⚠️",
            )
    else:
        st.caption("No replicated design points for this response — Lack of Fit could not be tested against Pure Error.")

    st.markdown("#### ANOVA Table")
    display_df = anova_df.copy()
    display_df["p-value (Prob > F)"] = display_df["p-value (Prob > F)"].apply(
        lambda v: "" if pd.isna(v) else ("< 0.0001" if v < 0.0001 else f"{v:.4f}")
    )
    display_df["F-value"] = display_df["F-value"].apply(lambda v: "" if pd.isna(v) else f"{v:.2f}")
    display_df["SS"] = display_df["SS"].apply(lambda v: "" if pd.isna(v) else f"{v:.4f}")
    display_df["MS"] = display_df["MS"].apply(lambda v: "" if pd.isna(v) else f"{v:.4f}")
    display_df["df"] = display_df["df"].apply(lambda v: "" if pd.isna(v) else f"{int(v)}")

    def highlight_totals(row):
        return ["font-weight: 600; background-color: #D6EFE9" if row["Source"] in ("Model", "Cor Total") else "" for _ in row]

    st.dataframe(
        display_df.style.apply(highlight_totals, axis=1),
        width="stretch", hide_index=True,
    )
    st.caption("Significance judged at α = 0.05 (p < 0.05 = significant).")

    st.markdown("#### Fit Statistics")
    m1, m2, m3, m4, m5, m6 = st.columns(6)
    m1.metric("R²", f"{fs['R²']:.4f}")
    m2.metric("Adjusted R²", f"{fs['Adjusted R²']:.4f}")
    m3.metric("Predicted R²", f"{fs['Predicted R²']:.4f}")
    m4.metric("Adeq Precision", f"{fs['Adequate Precision']:.3f}")
    m5.metric("C.V. %", f"{fs['C.V. %']:.2f}")
    m6.metric("Std. Dev.", f"{fs['Std. Dev.']:.4f}")

    if abs(fs["Adjusted R²"] - fs["Predicted R²"]) > 0.2:
        st.warning(
            "The difference between Adjusted R² and Predicted R² is more than "
            "0.2 — this can indicate a large block effect, a possible problem "
            "with the model or data, or simply the limits of a 10-run design. "
            "Report this caveat alongside the equation.",
            icon="⚠️",
        )
    if fs["Adequate Precision"] < 4:
        st.warning(
            "Adequate Precision is below 4, indicating an inadequate signal-to-"
            "noise ratio — this model should not be used to navigate the design "
            "space for this response.",
            icon="⚠️",
        )
    else:
        st.caption(f"Adequate Precision = {fs['Adequate Precision']:.3f} (> 4 is desirable — indicates an adequate signal).")

    if result["cond_no"] > 10:
        st.caption(
            f"Condition number (coded design) = {result['cond_no']:.1f} — "
            + ("above the common rule-of-thumb of 10, indicating multicollinearity between quadratic terms."
               if result["cond_no"] <= 30 else
               "well above 10, indicating substantial multicollinearity between quadratic terms; interpret individual coefficients with caution.")
        )

    st.markdown("#### Final Equation")
    tab_actual, tab_coded = st.tabs(["In Terms of Actual Factors", "In Terms of Coded Factors"])
    with tab_actual:
        st.code(result["eq_actual"], language=None)
        st.caption("Use actual-factor coefficients to predict the response for given levels of A and B in their original units.")
    with tab_coded:
        st.code(result["eq_coded"], language=None)
        st.caption(
            "Coded factors range from −1 (low level) to +1 (high level) over the "
            "explored design space, so coefficient magnitude directly reflects "
            "each term's relative effect on the response."
        )

    st.markdown("#### Export")
    docx_buf = build_anova_docx(response_name, result)
    st.download_button(
        f"Download ANOVA table (Word) — {response_name}",
        data=docx_buf,
        file_name=f"ANOVA_{response_name.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    st.caption(
        "The exported table matches this page (Source / SS / df / MS / F-value / "
        "p-value / Significance), followed by fit statistics and both final "
        "equations — formatted to drop directly into a manuscript's Results "
        "and Discussion section."
    )

# ---------------- Response surfaces ----------------
elif page == "Response Surfaces":
    render_header(
        "Response Surface Methodology",
        "3D Surface & Contour Plots",
        "Model-predicted response across the full Stearic acid × Tween 80 design space.",
    )

    c1, c2 = st.columns(2)
    surface_model_name = c1.selectbox(
        "Model",
        list(models.keys()),
        index=list(models.keys()).index(best_model_name),
        help="Defaults to the model with the best leave-one-out cross-validated R².",
    )
    target_options = {
        "Entrapment efficiency": 0,
        "Drug content": 1,
        "Drug release": 2,
        "Particle size": 3,
    }
    target_name = c2.selectbox("Response to plot", list(target_options.keys()))
    target_index = target_options[target_name]

    f1_range = np.linspace(X["Stearic acid"].min(), X["Stearic acid"].max(), 30)
    f2_range = np.linspace(X["Tween 80"].min(), X["Tween 80"].max(), 30)
    f1_grid, f2_grid = np.meshgrid(f1_range, f2_range)
    inputs = np.array([[a, b] for a, b in zip(np.ravel(f1_grid), np.ravel(f2_grid))])
    preds = models[surface_model_name].predict(inputs)[:, target_index].reshape(f1_grid.shape)

    tab_surface, tab_contour = st.tabs(["3D Surface", "2D Contour"])
    with tab_surface:
        fig3d = go.Figure(data=[go.Surface(
            x=f1_range, y=f2_range, z=preds, colorscale=TEAL_SCALE,
            colorbar=dict(title=target_name),
        )])
        fig3d.add_trace(go.Scatter3d(
            x=data["Stearic acid"], y=data["Tween 80"], z=data[target_name],
            mode="markers", marker=dict(size=5, color="#C97A2B"), name="Experimental runs",
        ))
        fig3d.update_layout(
            scene=dict(xaxis_title="Stearic acid (mg)", yaxis_title="Tween 80 (mg)", zaxis_title=target_name),
            height=560, margin=dict(l=0, r=0, t=20, b=0),
            paper_bgcolor="rgba(0,0,0,0)", font=dict(family="Inter, sans-serif", color="#101820"),
        )
        st.plotly_chart(fig3d, width='stretch')

    with tab_contour:
        figc = go.Figure(data=go.Contour(
            x=f1_range, y=f2_range, z=preds, colorscale=TEAL_SCALE,
            contours=dict(showlabels=True, labelfont=dict(size=10, color="white")),
            colorbar=dict(title=target_name),
        ))
        figc.add_trace(go.Scatter(
            x=data["Stearic acid"], y=data["Tween 80"], mode="markers",
            marker=dict(size=10, color="#C97A2B", symbol="x"), name="Experimental runs",
        ))
        figc.update_layout(
            xaxis_title="Stearic acid (mg)", yaxis_title="Tween 80 (mg)",
            height=520, margin=dict(l=10, r=10, t=20, b=10),
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            font=dict(family="Inter, sans-serif", color="#101820"),
        )
        st.plotly_chart(figc, width='stretch')

# ---------------- Optimization ----------------
elif page == "Optimization":
    render_header(
        "Desirability",
        "Formulation Optimization",
        "Best experimental run by desirability = (Entrapment efficiency × Drug content) / Particle size.",
    )
    desirability = (y["Entrapment efficiency"] * y["Drug content"]) / y["Particle size"]
    best_idx = desirability.idxmax()
    best_row = data.loc[best_idx]

    st.markdown("#### Best formulation (experimental validation)")
    m1, m2, m3, m4, m5, m6 = st.columns(6)
    m1.metric("Run", int(best_row["Runs"]))
    m2.metric("Stearic acid", f"{best_row['Stearic acid']:.0f} mg")
    m3.metric("Tween 80", f"{best_row['Tween 80']:.0f} mg")
    m4.metric("Entrapment eff.", f"{best_row['Entrapment efficiency']:.3f}")
    m5.metric("Drug content", f"{best_row['Drug content']:.3f}")
    m6.metric("Particle size", f"{best_row['Particle size']:.1f} nm")

    st.dataframe(data.loc[[best_idx]], width='stretch', hide_index=True)
    st.caption(
        "Desirability = (Entrapment efficiency × Drug content) / Particle size, "
        "evaluated on the actual experimental runs (not model-predicted)."
    )

# ---------------- Outlier Analysis ----------------
elif page == "Outlier Analysis":
    render_header("Quality Control", "Outlier Detection", "Boxplots and z-score screening across every variable.")

    st.markdown("#### Independent Variables")
    fig = make_subplots(rows=1, cols=2, subplot_titles=["Stearic acid", "Tween 80"])
    fig.add_trace(go.Box(y=data["Stearic acid"], name="Stearic acid", marker_color="#0E6E62", boxmean=True), row=1, col=1)
    fig.add_trace(go.Box(y=data["Tween 80"], name="Tween 80", marker_color="#0E6E62", boxmean=True), row=1, col=2)
    fig.update_layout(showlegend=False, height=380, margin=dict(l=10, r=10, t=40, b=10),
                       paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                       font=dict(family="Inter, sans-serif", color="#101820"))
    st.plotly_chart(fig, width='stretch')

    st.markdown("#### Dependent Variables")
    dep_cols = ["Entrapment efficiency", "Drug content", "Drug release", "Particle size"]
    fig2 = make_subplots(rows=1, cols=4, subplot_titles=dep_cols)
    for i, col in enumerate(dep_cols, start=1):
        fig2.add_trace(go.Box(y=data[col], name=col, marker_color="#C97A2B", boxmean=True), row=1, col=i)
    fig2.update_layout(showlegend=False, height=380, margin=dict(l=10, r=10, t=40, b=10),
                        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                        font=dict(family="Inter, sans-serif", color="#101820"))
    st.plotly_chart(fig2, width='stretch')

    st.markdown("#### Z-Score Outlier Detection")
    numeric_data = data.select_dtypes(include=[np.number])
    z_scores = np.abs(zscore(numeric_data))
    outliers = (z_scores > 2).any(axis=1)

    fig3 = go.Figure()
    fig3.add_trace(go.Scatter(
        x=list(range(len(data))), y=data["Entrapment efficiency"], mode="markers",
        marker=dict(size=12, color=np.where(outliers, "#C97A2B", "#0E6E62")),
        text=[f"Run {r}" for r in data["Runs"]], hovertemplate="%{text}<br>%{y}<extra></extra>",
    ))
    fig3.update_layout(
        xaxis_title="Run Index", yaxis_title="Entrapment efficiency",
        height=380, margin=dict(l=10, r=10, t=20, b=10),
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(family="Inter, sans-serif", color="#101820"),
    )
    st.plotly_chart(fig3, width='stretch')

    st.markdown("#### Outlier Runs (All Variables)")
    st.write(data[outliers])
